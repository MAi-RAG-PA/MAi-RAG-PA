#!/usr/bin/env python3
# ~/MAi-RAG/scripts/ingest_full_pipeline_with_progress.py

import os
import sys
import json
import hashlib
import uuid
import asyncio
import argparse
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import logging

# Core libs
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http.models import PointStruct, Distance, VectorParams

# Document parsing
import pdfplumber
import spacy
from bs4 import BeautifulSoup
from ebooklib import epub
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Install python-docx for DOCX support: pip install python-docx")

# Markdown/HTML
import frontmatter
from markdown_it import MarkdownIt

# Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Rich document metadata"""
    filename: str
    source_type: str
    author: Optional[str] = None
    created: Optional[str] = None
    modified: Optional[str] = None
    title: Optional[str] = None
    tags: List[str] = None
    version_hash: str = ""
    word_count: int = 0
    file_hash: str = ""

class ProgressCallback:
    """Async progress reporting"""
    def __init__(self):
        self.callback = None
    
    async def set_callback(self, cb):
        self.callback = cb
    
    async def __call__(self, stage: str, current: int, total: int):
        if self.callback:
            await self.callback(stage, current, total)

progress = ProgressCallback()

# Global models/clients
nlp = None
embedding_model = None
client = None

def init_models():
    global nlp, embedding_model, client
    nlp = spacy.load("en_core_web_sm")
    nlp.max_length = 2_000_000
    embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    client = QdrantClient("localhost", port=6333)

def file_hash(filepath: Path) -> str:
    """SHA256 hash for change detection"""
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            h.update(chunk)
    return h.hexdigest()

def extract_metadata(filepath: Path) -> DocumentMetadata:
    """Extract rich metadata from any format"""
    stat = filepath.stat()
    meta = DocumentMetadata(
        filename=filepath.name,
        source_type=filepath.suffix[1:],
        created=datetime.fromtimestamp(stat.st_ctime).isoformat(),
        modified=datetime.fromtimestamp(stat.st_mtime).isoformat(),
        file_hash=file_hash(filepath),
        tags=[]
    )
    
    # Format-specific metadata
    ext = meta.source_type.lower()
    if ext == 'pdf':
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        meta.title = doc.metadata.get('title', '')
        meta.author = doc.metadata.get('author', '')
        doc.close()
    elif ext == 'docx' and DOCX_AVAILABLE:
        doc = Document(filepath)
        meta.title = doc.core_properties.title or ''
        meta.author = doc.core_properties.author or ''
    elif ext == 'md':
        with open(filepath) as f:
            fm = frontmatter.load(f)
            meta.title = fm.get('title')
            meta.author = fm.get('author')
            meta.tags = fm.get('tags', [])
    
    # Auto-extract tags from content
    sample = open(filepath).read(2000)
    doc = nlp(sample)
    entities = [ent.text for ent in doc.ents if ent.label_ in ['PERSON', 'ORG', 'GPE']]
    meta.tags.extend(entities[:5])
    
    meta.version_hash = meta.file_hash  # Track changes
    return meta

def extract_text(filepath: Path) -> str:
    """Multi-format text extraction"""
    ext = filepath.suffix.lower()
    
    if ext == '.pdf':
        text = ''
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n\n'
        return text.strip()
    
    elif ext == '.docx' and DOCX_AVAILABLE:
        doc = Document(filepath)
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    
    elif ext == '.epub':
        book = epub.read_epub(str(filepath))
        text = ''
        for item in book.get_items():
            if item.get_type() == 1:  # ITEM_DOCUMENT
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                text += soup.get_text() + '\n\n'
        return text.strip()
    
    elif ext in ['.txt', '.md']:
        encoding = 'utf-8'
        try:
            with open(filepath, encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            with open(filepath, encoding='latin1') as f:
                return f.read()
    
    elif ext == '.html':
        with open(filepath) as f:
            soup = BeautifulSoup(f, 'html.parser')
            return soup.get_text()
    
    else:
        logger.warning(f"Unsupported: {filepath}")
        return ""

def smart_chunk(text: str, meta: DocumentMetadata, max_words=300, overlap=50) -> List[Dict[str, Any]]:
    """Semantic chunking with metadata preservation"""
    doc = nlp(text)
    chunks = []
    current_chunk = []
    current_words = 0
    
    for sent in doc.sents:
        sent_words = len(sent.text.split())
        if current_words + sent_words > max_words and current_chunk:
            chunks.append({
                'text': ' '.join(current_chunk),
                'metadata': meta.__dict__.copy(),
                'word_count': current_words
            })
            # Overlap
            current_chunk = current_chunk[-overlap//5:]
            current_words = sum(len(s.split()) for s in current_chunk)
        
        current_chunk.append(sent.text)
        current_words += sent_words
    
    if current_chunk:
        chunks.append({
            'text': ' '.join(current_chunk),
            'metadata': meta.__dict__.copy(),
            'word_count': current_words
        })
    
    return chunks

async def ingest_document(filepath: Path, collection_name: str, force=False):
    """Incremental ingestion with change detection"""
    meta_file = filepath.parent / f"{filepath.stem}.meta.json"
    
    # Check if needs reprocessing
    current_hash = file_hash(filepath)
    if meta_file.exists():
        with open(meta_file) as f:
            stored = json.load(f)
        if stored.get('file_hash') == current_hash and not force:
            logger.info(f"Skipping unchanged: {filepath}")
            return 0
    
    logger.info(f"Processing: {filepath}")
    
    # Extract metadata + text
    meta = extract_metadata(filepath)
    text = extract_text(filepath)
    if not text.strip():
        logger.warning(f"No text: {filepath}")
        return 0
    
    # Chunk the text
    chunks = smart_chunk(text, meta)
    
    # Embed and store
    points = []
    for i, chunk in enumerate(chunks):
        vector = embedding_model.encode(chunk['text']).tolist()
        point_id = str(uuid.uuid4())
        points.append(PointStruct(
            id=point_id,
            vector=vector,
            payload=chunk
        ))
    
    # Upsert to Qdrant
    if points:
        client.upsert(collection_name=collection_name, points=points)
    
    # Save metadata
    with open(meta_file, 'w') as f:
        json.dump(meta.__dict__, f, indent=2)
    
    return len(chunks)

async def ingest_directory(input_dir: str, chunks_dir: str, collection_name: str, max_words: int = 300, force_reingest: bool = False):
    """Ingest all documents from a directory"""
    init_models()
    
    # Ensure collection exists
    try:
        client.get_collection(collection_name)
    except:
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=384, distance=Distance.COSINE)
        )
    
    input_path = Path(input_dir).expanduser()
    
    # Find documents
    files = list(input_path.rglob("*"))
    doc_files = [f for f in files if f.is_file() and f.suffix.lower() in 
                 ['.pdf', '.docx', '.txt', '.md', '.html', '.epub']]
    
    print(f"📁 Found {len(doc_files)} documents")
    
    total_chunks = 0
    for i, filepath in enumerate(doc_files, 1):
        print(f"[{i}/{len(doc_files)}] {filepath.name}")
        chunks = await ingest_document(filepath, collection_name, force=force_reingest)
        total_chunks += chunks
        
        # Progress callback
        if progress.callback:
            await progress("processing", i, len(doc_files))
    
    # Final Qdrant upsert summary
    print(f"✅ Ingested {total_chunks} chunks → '{collection_name}'")

async def main():
    parser = argparse.ArgumentParser(description="Enhanced MAi-RAG Ingestion v2.0")
    parser.add_argument("--collection", help="Qdrant collection name", default="local_docs")
    parser.add_argument("--input-dir", default="~/MAi-RAG/documents_storage")
    parser.add_argument("--chunks-dir", default="~/MAi-RAG/chunks_output")
    parser.add_argument("--max-words", type=int, default=300)
    parser.add_argument("--force", action="store_true", help="Reingest all")
    
    args = parser.parse_args()
    
    await ingest_directory(
        args.input_dir,
        args.chunks_dir,
        args.collection,
        args.max_words,
        args.force
    )

if __name__ == "__main__":
    asyncio.run(main())